from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "pitch/MindGraph-Notes-One-Pager-Publikum.docx"
CONTENT_DXA = 10490

INK = RGBColor(28, 32, 36)
MUTED = RGBColor(91, 99, 109)
BLUE = RGBColor(24, 87, 139)
TEAL = RGBColor(16, 116, 109)
ORANGE = RGBColor(196, 101, 55)
GOLD = RGBColor(151, 111, 28)
PALE_BLUE = "EAF3FA"
PALE_TEAL = "E8F6F3"
PALE_ORANGE = "FFF0E7"
PALE_GOLD = "FFF7DA"
LINE = "D7DEE6"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_tbl_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_col_widths(table, widths):
    grid = table._tbl.tblGrid
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            set_cell_width(cell, width)


def set_row_height(row, height):
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(height))
    tr_pr.append(tr_h)


def set_font(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def para(cell_or_doc, text="", size=10.5, color=INK, bold=False, italic=False, before=0, after=4, align=None):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)


def add_badge(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_tbl_width(table, 3600)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_border(cell, "B9DDD7", "8")
    set_cell_margins(cell, 70, 140, 70, 140)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    set_font(r, size=8.5, color=TEAL, bold=True)


def add_icon_heading(cell, label, title, color, fill):
    clear_cell(cell)
    set_cell_shading(cell, fill)
    set_cell_border(cell, LINE, "8")
    set_cell_margins(cell, 130, 150, 120, 150)
    p = para(cell, label, size=8.5, color=color, bold=True, after=3)
    p.paragraph_format.line_spacing = 1
    para(cell, title, size=12.2, color=INK, bold=True, after=5)


def add_bullets(cell, items):
    for item in items:
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        marker = p.add_run("• ")
        set_font(marker, size=9.4, color=ORANGE, bold=True)
        r = p.add_run(item)
        set_font(r, size=9.4, color=INK)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = INK
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("MindGraph Notes | Pitchtag Mittelhessen")
    set_font(hr, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Download: mindgraph-notes.de  |  Open Source  |  Lokale KI ohne Cloud-Zwang")
    set_font(fr, size=8.6, color=MUTED)

    add_badge(doc, "One-Pager für Teilnehmer")
    para(doc, "MindGraph Notes", size=27, color=BLUE, bold=True, after=0)
    para(doc, "Zeigt dir, was heute wichtig ist.", size=17, color=ORANGE, bold=True, after=6)
    lead = para(
        doc,
        "Ein lokaler Arbeitsraum, der Notizen, Aufgaben, E-Mails und Dokumente verbindet - "
        "und morgens sichtbar macht, worauf es wirklich ankommt.",
        size=11.6,
        color=INK,
        after=8,
    )
    lead.paragraph_format.line_spacing = 1.12

    hero = doc.add_table(rows=1, cols=2)
    set_tbl_width(hero, CONTENT_DXA)
    set_col_widths(hero, [6100, 4390])
    left, right = hero.rows[0].cells
    set_cell_shading(left, PALE_BLUE)
    set_cell_border(left, "C8DAEA", "8")
    set_cell_margins(left, 155, 170, 145, 170)
    clear_cell(left)
    para(left, "Das Problem im Büro", size=8.7, color=BLUE, bold=True, after=3)
    para(left, "Information ist da - aber nicht sortiert.", size=14.2, color=INK, bold=True, after=5)
    para(
        left,
        "Mails, Termine, Projektakten, PDFs und Aufgaben liegen verstreut. Man verbringt Zeit mit Suchen, "
        "Nachhalten und Erinnern statt mit Entscheiden.",
        size=9.8,
        color=INK,
        after=2,
    )
    set_cell_shading(right, PALE_ORANGE)
    set_cell_border(right, "E7C8B6", "8")
    set_cell_margins(right, 155, 170, 145, 170)
    clear_cell(right)
    para(right, "Der Nutzen", size=8.7, color=ORANGE, bold=True, after=3)
    para(right, "Ein Tagesblick statt fünf Suchfenster.", size=14.2, color=INK, bold=True, after=5)
    para(
        right,
        "MindGraph Notes bündelt relevante Hinweise, Aufgaben und offene Antworten - lokal auf dem eigenen Rechner.",
        size=9.8,
        color=INK,
        after=2,
    )

    para(doc, "", after=1)
    grid = doc.add_table(rows=1, cols=3)
    set_tbl_width(grid, CONTENT_DXA)
    set_col_widths(grid, [3497, 3497, 3496])
    cards = [
        ("1", "Morning Briefing", BLUE, PALE_BLUE, [
            "Tagesüberblick aus Notizen, Aufgaben und E-Mails",
            "Relevanz statt chronologischer Flut",
            "Offene Antworten und Fristen im Blick",
        ]),
        ("2", "Arbeitswissen bleibt verbunden", TEAL, PALE_TEAL, [
            "Markdown-Notizen mit Wiki-Links und Backlinks",
            "Projektakten, Meetings und Dokumente an einem Ort",
            "Canvas für Zusammenhänge, wenn es komplex wird",
        ]),
        ("3", "KI, die vertraulich bleibt", ORANGE, PALE_ORANGE, [
            "Lokale Modelle über Ollama oder LM Studio",
            "Keine Telemetrie, kein Cloud-Zwang",
            "E2E-verschlüsselter Sync bei Bedarf",
        ]),
    ]
    for idx, cell in enumerate(grid.rows[0].cells):
        label, title, color, fill, items = cards[idx]
        add_icon_heading(cell, label, title, color, fill)
        add_bullets(cell, items)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    para(doc, "", after=0)
    two = doc.add_table(rows=1, cols=2)
    set_tbl_width(two, CONTENT_DXA)
    set_col_widths(two, [6000, 4490])
    workflow, invite = two.rows[0].cells
    set_cell_shading(workflow, WHITE)
    set_cell_border(workflow, LINE, "8")
    set_cell_margins(workflow, 120, 160, 120, 160)
    clear_cell(workflow)
    para(workflow, "Typischer Ablauf", size=8.8, color=BLUE, bold=True, after=3)
    steps = [
        ("1.", "E-Mail oder Dokument landet im Arbeitsraum."),
        ("2.", "MindGraph Notes erkennt Aufgaben, Termine und Kontext."),
        ("3.", "Der Tagesblick zeigt, was wirklich nach Aufmerksamkeit verlangt."),
    ]
    for n, text in steps:
        p = workflow.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        nr = p.add_run(n + " ")
        set_font(nr, size=10.5, color=ORANGE, bold=True)
        tr = p.add_run(text)
        set_font(tr, size=9.8, color=INK)

    set_cell_shading(invite, PALE_GOLD)
    set_cell_border(invite, "E5D18A", "8")
    set_cell_margins(invite, 120, 160, 120, 160)
    clear_cell(invite)
    para(invite, "Gesucht: 10 Test-User aus Mittelhessen", size=8.8, color=GOLD, bold=True, after=3)
    para(invite, "Passt besonders, wenn Sie ...", size=11.2, color=INK, bold=True, after=3)
    add_bullets(invite, [
        "viele Mails, Termine und Projektinfos koordinieren",
        "sensibles Wissen lieber lokal behalten",
        "KI praktisch testen wollen, ohne gleich die Organisation umzubauen",
    ])

    cta = doc.add_table(rows=1, cols=3)
    set_tbl_width(cta, CONTENT_DXA)
    set_col_widths(cta, [3300, 4190, 3000])
    cta_items = [
        ("Website", "mindgraph-notes.de"),
        ("Code", "github.com/bydb/mindgraph-notes"),
        ("Kontakt", "Jochen Leeder | bydb.io"),
    ]
    for i, cell in enumerate(cta.rows[0].cells):
        clear_cell(cell)
        set_cell_shading(cell, "F6F8FA")
        set_cell_border(cell, LINE, "8")
        set_cell_margins(cell, 80, 130, 80, 130)
        para(cell, cta_items[i][0], size=8, color=MUTED, bold=True, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        para(cell, cta_items[i][1], size=9.4, color=BLUE, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_together = True

    doc.save(OUT)


if __name__ == "__main__":
    build()
